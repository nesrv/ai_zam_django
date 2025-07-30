import json
import logging
from django.http import JsonResponse, FileResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.views.decorators.clickjacking import xframe_options_exempt
from django.shortcuts import render
from django.db.models import Count, Q
from django.utils import timezone
from datetime import timedelta
from .services import process_telegram_update, send_telegram_message, check_bot_token, BOT_TOKEN
from .models import TelegramUser, TelegramMessage
import io
from fpdf import FPDF
from docx import Document
import pandas as pd
from telegrambot.models import TemporaryDocument
import os
from io import BytesIO
from fpdf.errors import FPDFException

logger = logging.getLogger(__name__)

@csrf_exempt
@require_POST
def telegram_webhook(request):
    print('=== ВЫЗВАН telegram_webhook ===')
    import logging
    logging.getLogger('django').warning('=== ВЫЗВАН telegram_webhook ===')
    """Webhook для получения сообщений от Telegram"""
    logger.info("=== WEBHOOK ЗАПРОС ===")
    
    try:
        # Логируем базовую информацию о запросе
        logger.info(f"Метод: {request.method}")
        logger.info(f"Content-Type: {request.content_type}")
        logger.info(f"User-Agent: {request.META.get('HTTP_USER_AGENT', 'Не указан')}")
        logger.info(f"Remote IP: {request.META.get('REMOTE_ADDR', 'Неизвестно')}")
        logger.info(f"Размер тела запроса: {len(request.body)} байт")
        
        if request.method != 'POST':
            logger.warning(f"Неверный метод: {request.method}")
            return JsonResponse({'error': 'Method not allowed'}, status=405)
        
        # Проверяем Content-Type
        content_type = request.content_type or ''
        if 'application/json' not in content_type:
            logger.warning(f"Неверный Content-Type: {content_type}")
            return JsonResponse({'error': 'Content-Type must be application/json'}, status=400)
        
        # Проверяем размер тела запроса
        if len(request.body) == 0:
            logger.warning("Пустое тело запроса")
            return JsonResponse({'error': 'Empty request body'}, status=400)
        
        # Проверяем User-Agent
        user_agent = request.META.get('HTTP_USER_AGENT', '')
        remote_addr = request.META.get('REMOTE_ADDR', '')
        
        # Простая проверка IP (можно улучшить с помощью ipaddress модуля)
        if 'TelegramBot' not in user_agent and 'python-telegram-bot' not in user_agent:
            logger.warning(f"Подозрительный запрос к webhook от: {remote_addr} - {user_agent}")
            # Не блокируем, но логируем для безопасности
        
        # Декодируем JSON
        try:
            update_data = json.loads(request.body.decode('utf-8'))
            logger.info(f"Успешно декодирован JSON: {update_data}")
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка декодирования JSON: {e}")
            logger.error(f"Тело запроса: {request.body}")
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            logger.error(f"Неожиданная ошибка при декодировании JSON: {e}")
            return JsonResponse({'error': 'JSON processing error'}, status=400)
        
        # Проверяем структуру update
        if not isinstance(update_data, dict):
            logger.error(f"Update не является словарем: {type(update_data)}")
            return JsonResponse({'error': 'Invalid update format'}, status=400)
        
        if 'update_id' not in update_data:
            logger.error("Отсутствует update_id в update")
            return JsonResponse({'error': 'Missing update_id'}, status=400)
        
        # Обрабатываем сообщения из всех типов чатов
        message = update_data.get('message', {})
        if message:
            # Используем новую функцию для обработки сообщений
            from .services import process_telegram_message
            process_telegram_message(message)
            logger.info(f"Сообщение обработано и сохранено в базу данных")
        
        # Обрабатываем update
        try:
            result = process_telegram_update(update_data)
            logger.info(f"Update обработан успешно: {result}")
            return JsonResponse({'status': 'ok'})
        except Exception as e:
            logger.error(f"Ошибка при обработке update: {e}")
            return JsonResponse({'error': 'Update processing error'}, status=500)
            
    except Exception as e:
        logger.error(f"Критическая ошибка в webhook: {e}")
        logger.error(f"Тип ошибки: {type(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return JsonResponse({'error': 'Internal server error'}, status=500)

def bot_status(request):
    """Страница статуса бота с чатом"""
    try:
        # Получаем статистику
        total_users = TelegramUser.objects.count()
        active_users = TelegramUser.objects.filter(is_active=True).count()
        total_messages = TelegramMessage.objects.count()
        
        # Получаем все сообщения для чата, отсортированные по времени
        all_messages = list(TelegramMessage.objects.select_related('user').order_by('created_at'))
        
        # Получаем последние 5 сообщений из AI чата
        try:
            from ai.models import ChatMessage
            ai_messages = ChatMessage.objects.select_related('session').order_by('-created_at')[:5]
            
            # Добавляем AI сообщения в правильном порядке
            for msg in reversed(ai_messages):
                # Обрезаем длинные сообщения
                content = msg.content
                if len(content) > 200:
                    content = content[:200] + '...'
                
                # Создаем объект похожий на TelegramMessage
                ai_wrapper = type('AIMessage', (), {
                    'content': content,
                    'created_at': msg.created_at,
                    'is_from_user': msg.message_type == 'user',
                    'file': getattr(msg, 'file', None),
                    'user': type('User', (), {
                        'first_name': 'AI User' if msg.message_type == 'user' else 'DeepSeek'
                    })()
                })()
                
                all_messages.insert(0, ai_wrapper)
                
        except Exception as e:
            logger.error(f"Ошибка получения AI сообщений: {e}")
        
        # Последние сообщения для статистики (первые 10)
        recent_messages = all_messages[-10:] if all_messages else []
        
        context = {
            'total_users': total_users,
            'active_users': active_users,
            'total_messages': total_messages,
            'recent_messages': recent_messages,
            'all_messages': all_messages,  # Все сообщения для чата
        }
        
        return render(request, 'telegrambot/status.html', context)
        
    except Exception as e:
        logger.error(f"Ошибка получения статуса бота: {e}")
        return JsonResponse({'error': str(e)}, status=500)

def bot_management(request):
    """Страница управления ботом"""
    try:
        # Получаем список пользователей
        users = TelegramUser.objects.order_by('-created_at')[:50]
        
        context = {
            'users': users,
        }
        
        return render(request, 'telegrambot/management.html', context)
        
    except Exception as e:
        logger.error(f"Ошибка управления ботом: {e}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def send_broadcast(request):
    """Отправка массового сообщения всем пользователям"""
    try:
        data = json.loads(request.body)
        message = data.get('message', '').strip()
        
        if not message:
            return JsonResponse({'error': 'Сообщение не может быть пустым'}, status=400)
        
        # Получаем всех активных пользователей
        users = TelegramUser.objects.filter(is_active=True)
        total_users = users.count()
        
        if total_users == 0:
            return JsonResponse({
                'ok': False,
                'error': 'Нет активных пользователей для отправки сообщения',
                'sent_count': 0,
                'total_users': 0
            })
        
        # Проверяем токен бота перед отправкой
        token = BOT_TOKEN
        is_valid, bot_info = check_bot_token(token)
        if not is_valid:
            return JsonResponse({
                'ok': False,
                'error': f'Недействительный токен бота: {bot_info}',
                'sent_count': 0,
                'total_users': total_users
            })
        
        sent_count = 0
        failed_count = 0
        errors = []
        
        for user in users:
            try:
                logger.info(f"Отправка сообщения пользователю {user.telegram_id} ({user.first_name})")
                result = send_telegram_message(user.telegram_id, message)
                
                if result and result.get('ok'):
                    sent_count += 1
                    logger.info(f"✅ Сообщение успешно отправлено пользователю {user.telegram_id}")
                else:
                    failed_count += 1
                    error_msg = f"Ошибка отправки пользователю {user.telegram_id}: {result}"
                    errors.append(error_msg)
                    logger.error(error_msg)
                    
                    # Если пользователь заблокировал бота, помечаем его как неактивного
                    if result and result.get('error_code') == 400:
                        description = result.get('description', '').lower()
                        if 'bot was blocked' in description or 'chat not found' in description:
                            user.is_active = False
                            user.save()
                            logger.info(f"Пользователь {user.telegram_id} помечен как неактивный")
                    
            except Exception as e:
                failed_count += 1
                error_msg = f"Исключение при отправке пользователю {user.telegram_id}: {str(e)}"
                errors.append(error_msg)
                logger.error(error_msg)
        
        logger.info(f"Массовая рассылка завершена: {sent_count}/{total_users} успешно отправлено")
        
        return JsonResponse({
            'ok': True,
            'sent_count': sent_count,
            'failed_count': failed_count,
            'total_users': total_users,
            'errors': errors[:5] if errors else []  # Возвращаем только первые 5 ошибок
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        logger.error(f"Общая ошибка массовой рассылки: {e}")
        return JsonResponse({'error': str(e)}, status=500)

def webhook_status(request):
    """Страница статуса webhook для проверки"""
    if request.method == 'GET':
        return JsonResponse({
            'status': 'ok',
            'message': 'Telegram webhook endpoint is working',
            'method': 'GET',
            'note': 'This endpoint accepts only POST requests from Telegram'
        })
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)

def bot_dashboard(request):
    """Главная страница с историей и активностью бота"""
    try:
        # Общая статистика
        total_users = TelegramUser.objects.count()
        active_users = TelegramUser.objects.filter(is_active=True).count()
        total_messages = TelegramMessage.objects.count()
        
        # Статистика за последние 7 дней
        week_ago = timezone.now() - timedelta(days=7)
        new_users_week = TelegramUser.objects.filter(created_at__gte=week_ago).count()
        messages_week = TelegramMessage.objects.filter(created_at__gte=week_ago).count()
        
        # Статистика за последние 24 часа
        day_ago = timezone.now() - timedelta(days=1)
        new_users_day = TelegramUser.objects.filter(created_at__gte=day_ago).count()
        messages_day = TelegramMessage.objects.filter(created_at__gte=day_ago).count()
        
        # Последние активные пользователи
        recent_active_users = TelegramUser.objects.filter(
            messages__created_at__gte=day_ago
        ).distinct().order_by('-updated_at')[:10]
        
        # Последние сообщения
        recent_messages = TelegramMessage.objects.select_related('user').order_by('-created_at')[:20]
        
        # Статистика по типам сообщений
        message_types = TelegramMessage.objects.values('message_type').annotate(
            count=Count('id')
        ).order_by('-count')
        
        # Топ пользователей по активности
        top_users = TelegramUser.objects.annotate(
            message_count=Count('messages')
        ).filter(message_count__gt=0).order_by('-message_count')[:10]
        
        # Активность по часам (последние 24 часа)
        hourly_activity = []
        for i in range(24):
            hour_start = timezone.now() - timedelta(hours=23-i)
            hour_end = hour_start + timedelta(hours=1)
            count = TelegramMessage.objects.filter(
                created_at__gte=hour_start,
                created_at__lt=hour_end
            ).count()
            hourly_activity.append({
                'hour': hour_start.hour,
                'count': count
            })
        
        context = {
            # Общая статистика
            'total_users': total_users,
            'active_users': active_users,
            'total_messages': total_messages,
            
            # Статистика за неделю
            'new_users_week': new_users_week,
            'messages_week': messages_week,
            
            # Статистика за день
            'new_users_day': new_users_day,
            'messages_day': messages_day,
            
            # Активность
            'recent_active_users': recent_active_users,
            'recent_messages': recent_messages,
            'message_types': message_types,
            'top_users': top_users,
            'hourly_activity': hourly_activity,
            
            # Статус бота
            'bot_status': 'active',  # Можно добавить реальную проверку статуса
        }
        
        return render(request, 'telegrambot/dashboard.html', context)
        
    except Exception as e:
        logger.error(f"Ошибка получения данных dashboard: {e}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def generate_document(request):
    """Генерация документа с помощью DeepSeek API"""
    try:
        data = json.loads(request.body)
        prompt = data.get('prompt', '').strip()
        
        logger.info(f"Получен запрос на генерацию документа: {prompt}")
        
        if not prompt:
            logger.error("Пустой запрос на генерацию документа")
            return JsonResponse({'error': 'Запрос не может быть пустым'}, status=400)
        
        # Импортируем функцию генерации документа
        from .services import generate_document_with_deepseek
        
        # Генерируем документ с помощью DeepSeek
        logger.info("Отправляю запрос к DeepSeek API...")
        generated_content = generate_document_with_deepseek(prompt)
        logger.info(f"Получен ответ от DeepSeek: {generated_content[:200]}...")
        
        if generated_content.startswith('Ошибка'):
            logger.error(f"Ошибка DeepSeek API: {generated_content}")
            return JsonResponse({
                'ok': False,
                'error': generated_content
            })
        
        # Получаем всех активных пользователей для отправки
        users = TelegramUser.objects.filter(is_active=True)
        total_users = users.count()
        logger.info(f"Найдено активных пользователей: {total_users}")
        
        if total_users == 0:
            logger.warning("Нет активных пользователей для отправки документа")
            return JsonResponse({
                'ok': False,
                'error': 'Нет активных пользователей для отправки документа',
                'generated_content': generated_content
            })
        
        # Отправляем сгенерированный документ всем пользователям
        sent_count = 0
        failed_count = 0
        
        for user in users:
            try:
                logger.info(f"Отправляю документ пользователю {user.telegram_id} ({user.first_name})")
                result = send_telegram_message(user.telegram_id, generated_content)
                
                if result and result.get('ok'):
                    sent_count += 1
                    logger.info(f"✅ Документ успешно отправлен пользователю {user.telegram_id}")
                else:
                    failed_count += 1
                    logger.error(f"Ошибка отправки документа пользователю {user.telegram_id}: {result}")
                    
                    # Если пользователь заблокировал бота, помечаем его как неактивного
                    if result and result.get('error_code') == 400:
                        description = result.get('description', '').lower()
                        if 'bot was blocked' in description or 'chat not found' in description:
                            user.is_active = False
                            user.save()
                            logger.info(f"Пользователь {user.telegram_id} помечен как неактивный")
                    
            except Exception as e:
                failed_count += 1
                logger.error(f"Исключение при отправке документа пользователю {user.telegram_id}: {str(e)}")
        
        logger.info(f"Документ сгенерирован и отправлен: {sent_count}/{total_users} пользователей")
        
        return JsonResponse({
            'ok': True,
            'generated_content': generated_content,
            'sent_count': sent_count,
            'failed_count': failed_count,
            'total_users': total_users
        })
        
    except json.JSONDecodeError as e:
        logger.error(f"Ошибка декодирования JSON: {e}")
        logger.error(f"Тело запроса: {request.body}")
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        logger.error(f"Ошибка генерации документа: {e}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
def clear_cache_view(request):
    """Временный view для очистки кэша (удалить после использования)"""
    if request.method == 'POST':
        try:
            from django.core.cache import cache
            cache.clear()
            
            # Пересобираем статические файлы
            from django.core.management import call_command
            call_command('collectstatic', '--noinput', '--clear')
            
            return JsonResponse({
                'status': 'success',
                'message': 'Кэш очищен и статические файлы пересобраны'
            })
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Ошибка: {str(e)}'
            }, status=500)
    
    return JsonResponse({
        'status': 'info',
        'message': 'Отправьте POST запрос для очистки кэша'
    })

@csrf_exempt
def export_document(request):
    """Экспорт сгенерированного документа в docx/pdf/xls (GET и POST)"""
    if request.method == 'POST':
        import json
        data = json.loads(request.body)
        doc_id = data.get('id')
        file_format = data.get('format', '').lower()
        content = None
        if doc_id:
            try:
                temp_doc = TemporaryDocument.objects.get(id=doc_id)
                content = temp_doc.content
            except TemporaryDocument.DoesNotExist:
                return JsonResponse({'error': 'Документ не найден'}, status=404)
        else:
            content = data.get('content', '').strip()
    elif request.method == 'GET':
        doc_id = request.GET.get('id')
        file_format = request.GET.get('format', '').lower()
        content = None
        if doc_id:
            try:
                temp_doc = TemporaryDocument.objects.get(id=doc_id)
                content = temp_doc.content
            except TemporaryDocument.DoesNotExist:
                return JsonResponse({'error': 'Документ не найден'}, status=404)
        else:
            content = request.GET.get('content', '').strip()
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    if not content or file_format not in ['docx', 'pdf', 'xls']:
        return JsonResponse({'error': 'Передайте id и format (docx/pdf/xls)'}, status=400)
    filename = f"document.{file_format}"
    if file_format == 'docx':
        doc = Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return FileResponse(buf, as_attachment=True, filename=filename)
    elif file_format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        font_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'fonts', 'DejaVuSans.ttf')
        pdf.add_font('DejaVu', '', font_path, uni=True)
        pdf.set_font('DejaVu', '', 10)
        content = content.replace('🔹', '-')

        lines = content.split('\n')
        table_started = False
        table_data = []
        for line in lines:
            # Поиск начала таблицы
            if line.strip().startswith('|') and line.strip().endswith('|'):
                table_started = True
                table_data.append([cell.strip() for cell in line.strip('|').split('|')])
            elif table_started and (line.strip().startswith('|') and line.strip().endswith('|')):
                table_data.append([cell.strip() for cell in line.strip('|').split('|')])
            elif table_started and not (line.strip().startswith('|') and line.strip().endswith('|')):
                # Таблица закончилась, выводим её
                if table_data:
                    # Определяем количество столбцов по первой строке
                    num_cols = len(table_data[0]) if table_data else 5
                    col_width = (pdf.w - pdf.l_margin - pdf.r_margin) / num_cols
                    for row in table_data:
                        for i, cell in enumerate(row):
                            if i < num_cols:  # Проверяем индекс
                                pdf.cell(col_width, 8, cell, border=1)
                        pdf.ln(8)
                table_started = False
                table_data = []
                if line.strip():
                    pdf.ln(4)
                    pdf.multi_cell(0, 8, line)
            else:
                if line.strip():
                    safe_multicell(pdf, 0, 8, line, max_len=80)
                else:
                    pdf.ln(4)
        # Если таблица была в самом конце
        if table_data:
            num_cols = len(table_data[0]) if table_data else 5
            col_width = (pdf.w - pdf.l_margin - pdf.r_margin) / num_cols
            for row in table_data:
                for i, cell in enumerate(row):
                    if i < num_cols:  # Проверяем индекс
                        pdf.cell(col_width, 8, cell, border=1)
                pdf.ln(8)

        pdf_output = BytesIO()
        pdf.output(pdf_output)
        pdf_output.seek(0)
        response = HttpResponse(pdf_output, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="document.pdf"'
        return response
    elif file_format == 'xls':
        lines = [l for l in content.split('\n') if l.strip()]
        df = pd.DataFrame({'Документ': lines})
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False)
        buf.seek(0)
        return FileResponse(buf, as_attachment=True, filename=filename)
    else:
        return JsonResponse({'error': 'Неподдерживаемый формат'}, status=400)

def safe_multicell(pdf, width, height, text, max_len=80):
    # Если ширина 0, используем ширину страницы минус отступы
    if width == 0:
        width = pdf.w - pdf.l_margin - pdf.r_margin
    
    # Проверяем, что ширина достаточна для одного символа
    if width < 10:
        width = 100
    
    try:
        words = text.split(' ')
        current_line = ''
        for word in words:
            # Если слово слишком длинное — разбиваем его
            while len(word) > max_len:
                part = word[:max_len]
                word = word[max_len:]
                if current_line:
                    pdf.multi_cell(width, height, current_line)
                    current_line = ''
                pdf.multi_cell(width, height, part)
            # Если добавление слова превышает лимит — печатаем текущую строку
            if len(current_line) + len(word) + 1 > max_len:
                if current_line:
                    pdf.multi_cell(width, height, current_line)
                current_line = word
            else:
                if current_line:
                    current_line += ' ' + word
                else:
                    current_line = word
        if current_line:
            pdf.multi_cell(width, height, current_line)
    except FPDFException as e:
        # Если возникла ошибка, просто добавляем текст как обычную строку
        pdf.ln(height)
        pdf.cell(0, height, text[:50] + '...' if len(text) > 50 else text)

@csrf_exempt
@require_POST
def create_object_ai(request):
    """Создание объекта с AI анализом последнего сообщения"""
    try:
        # Получаем последнее сообщение из чата
        last_message = TelegramMessage.objects.order_by('-created_at').first()
        
        if not last_message:
            return JsonResponse({
                'ok': False,
                'error': 'Нет сообщений в чате для анализа'
            })
        
        # Анализируем сообщение с помощью AI
        from .services import analyze_message_for_object_creation
        
        analysis_result = analyze_message_for_object_creation(last_message.content)
        
        if analysis_result.get('error'):
            return JsonResponse({
                'ok': False,
                'error': analysis_result['error']
            })
        
        # Возвращаем данные для создания объекта
        return JsonResponse({
            'ok': True,
            'object_data': analysis_result,
            'redirect_url': '/objects/create/',
            'message': 'Данные извлечены из последнего сообщения'
        })
        
    except Exception as e:
        logger.error(f"Ошибка создания объекта с AI: {e}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def send_file_to_deepseek(request):
    """Отправка файла и сообщения в DeepSeek для анализа"""
    try:
        message = request.POST.get('message', '').strip()
        uploaded_file = request.FILES.get('file')
        
        if not message and not uploaded_file:
            return JsonResponse({'error': 'Необходимо указать сообщение или прикрепить файл'}, status=400)
        
        # Формируем промпт для DeepSeek
        prompt = message or 'Проанализируй прикрепленный файл'
        
        if uploaded_file:
            prompt += f"\n\nПрикреплен файл: {uploaded_file.name} ({uploaded_file.size} байт)"
            
            # Извлекаем текст из файла в зависимости от типа
            if uploaded_file.content_type.startswith('text/') or uploaded_file.name.endswith(('.txt', '.md', '.py', '.js', '.html', '.css')):
                try:
                    file_content = uploaded_file.read().decode('utf-8')
                    prompt += f"\n\nСодержимое файла:\n{file_content[:2000]}"
                    if len(file_content) > 2000:
                        prompt += "\n\n[Файл обрезан для экономии токенов]"
                except UnicodeDecodeError:
                    prompt += "\n\n[Не удалось прочитать содержимое файла]"
            elif uploaded_file.name.endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(uploaded_file)
                    text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                    prompt += f"\n\nСодержимое DOCX файла:\n{text_content[:2000]}"
                    if len(text_content) > 2000:
                        prompt += "\n\n[Файл обрезан для экономии токенов]"
                except Exception as e:
                    prompt += f"\n\n[Ошибка чтения DOCX файла: {str(e)}]"
            elif uploaded_file.name.endswith('.pdf'):
                try:
                    import PyPDF2
                    from io import BytesIO
                    pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
                    text_content = ''
                    for page in pdf_reader.pages[:5]:  # Читаем только первые 5 страниц
                        text_content += page.extract_text() + '\n'
                    prompt += f"\n\nСодержимое PDF файла:\n{text_content[:2000]}"
                    if len(text_content) > 2000:
                        prompt += "\n\n[Файл обрезан для экономии токенов]"
                except Exception as e:
                    prompt += f"\n\n[Ошибка чтения PDF файла: {str(e)}]"
            elif uploaded_file.name.endswith(('.xlsx', '.xls')):
                try:
                    import pandas as pd
                    # Сохраняем файл в media/documents_ai/
                    from django.core.files.storage import default_storage
                    from django.core.files.base import ContentFile
                    import uuid
                    
                    # Сохраняем копию файла для чтения
                    file_content = uploaded_file.read()
                    file_name = f"{uuid.uuid4()}_{uploaded_file.name}"
                    file_path = default_storage.save(f'documents_ai/{file_name}', ContentFile(file_content))
                    
                    # Читаем Excel файл из байтов
                    from io import BytesIO
                    excel_buffer = BytesIO(file_content)
                    
                    engine = 'openpyxl' if uploaded_file.name.endswith('.xlsx') else 'xlrd'
                    
                    # Пытаемся прочитать все листы
                    try:
                        # Сначала получаем список листов
                        excel_file = pd.ExcelFile(excel_buffer, engine=engine)
                        sheet_names = excel_file.sheet_names
                        
                        text_content = f"📈 Анализ Excel файла: {uploaded_file.name}\n"
                        text_content += f"📄 Количество листов: {len(sheet_names)}\n"
                        text_content += f"📝 Названия листов: {', '.join(sheet_names)}\n\n"
                        
                        # Читаем первые 2 листа (или все, если их меньше)
                        sheets_to_read = sheet_names[:2]
                        
                        for sheet_name in sheets_to_read:
                            try:
                                df = pd.read_excel(excel_buffer, sheet_name=sheet_name, engine=engine, nrows=50)
                                
                                text_content += f"📉 Лист: '{sheet_name}'\n"
                                text_content += f"• Размер: {len(df)} строк × {len(df.columns)} столбцов\n"
                                text_content += f"• Столбцы: {', '.join(df.columns.astype(str))}\n"
                                
                                # Добавляем первые несколько строк данных
                                if not df.empty:
                                    # Показываем первые 5 строк
                                    sample_data = df.head(5).to_string(max_cols=8, max_colwidth=20)
                                    text_content += f"• Пример данных:\n{sample_data}\n"
                                    
                                    # Анализируем числовые столбцы
                                    numeric_cols = df.select_dtypes(include=['number']).columns
                                    if len(numeric_cols) > 0:
                                        text_content += f"• Числовые столбцы: {', '.join(numeric_cols)}\n"
                                        
                                        # Показываем статистику по первому числовому столбцу
                                        first_numeric = numeric_cols[0]
                                        stats = df[first_numeric].describe()
                                        text_content += f"• Статистика '{first_numeric}': мин={stats['min']:.2f}, макс={stats['max']:.2f}, среднее={stats['mean']:.2f}\n"
                                
                                text_content += "\n"
                                
                                # Пересоздаем buffer для следующего листа
                                excel_buffer = BytesIO(file_content)
                                
                            except Exception as sheet_error:
                                text_content += f"• Ошибка чтения листа '{sheet_name}': {str(sheet_error)}\n\n"
                                continue
                        
                    except Exception as e:
                        # Если не удалось прочитать как Excel файл, пробуем простое чтение
                        excel_buffer = BytesIO(file_content)
                        df = pd.read_excel(excel_buffer, engine=engine, nrows=50)
                        
                        text_content = f"📈 Анализ Excel файла: {uploaded_file.name}\n"
                        text_content += f"• Размер: {len(df)} строк × {len(df.columns)} столбцов\n"
                        text_content += f"• Столбцы: {', '.join(df.columns.astype(str))}\n\n"
                        text_content += df.to_string(max_rows=20, max_cols=8)
                    
                    prompt += f"\n\n{text_content[:3000]}"
                    if len(text_content) > 3000:
                        prompt += "\n\n[Данные обрезаны для экономии токенов]"
                        
                    # Сохраняем информацию о файле в ai_chatmessage
                    from ai.models import ChatSession, ChatMessage
                    session, created = ChatSession.objects.get_or_create(
                        session_id='telegram_excel',
                        defaults={'session_id': 'telegram_excel'}
                    )
                    
                    ChatMessage.objects.create(
                        session=session,
                        message_type='user',
                        content=f'Загружен Excel файл: {uploaded_file.name}',
                        file=file_path
                    )
                    
                except Exception as e:
                    logger.error(f"Ошибка обработки Excel файла: {str(e)}")
                    prompt += f"\n\n[Ошибка чтения Excel файла: {str(e)}]"
            else:
                prompt += "\n\n[Прикреплен файл - анализ содержимого недоступен для данного формата]"
        
        logger.info(f"Отправляю в DeepSeek промпт: {prompt[:200]}...")
        
        # Отправляем в DeepSeek
        from .services import generate_document_with_deepseek
        
        generated_content = generate_document_with_deepseek(prompt)
        
        if generated_content.startswith('Ошибка'):
            logger.error(f"Ошибка DeepSeek API: {generated_content}")
            return JsonResponse({
                'ok': False,
                'error': generated_content
            })
        
        logger.info(f"Получен ответ от DeepSeek: {generated_content[:200]}...")
        
        # Сохраняем ответ AI в базу данных
        try:
            from ai.models import ChatSession, ChatMessage
            session, created = ChatSession.objects.get_or_create(
                session_id='telegram_excel',
                defaults={'session_id': 'telegram_excel'}
            )
            
            ChatMessage.objects.create(
                session=session,
                message_type='assistant',
                content=generated_content
            )
        except Exception as e:
            logger.error(f"Ошибка сохранения ответа AI: {e}")
        
        return JsonResponse({
            'ok': True,
            'generated_content': generated_content,
            'message': 'Файл успешно проанализирован DeepSeek AI'
        })
        
    except Exception as e:
        logger.error(f"Ошибка отправки файла в DeepSeek: {e}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def create_object_from_message(request):
    """Создание объекта из содержимого сообщения"""
    try:
        data = json.loads(request.body)
        content = data.get('content', '').strip()
        
        if not content:
            return JsonResponse({
                'ok': False,
                'error': 'Содержимое сообщения не может быть пустым'
            })
        
        # Анализируем содержимое с помощью AI
        from .services import analyze_message_for_object_creation
        
        analysis_result = analyze_message_for_object_creation(content)
        
        if analysis_result.get('error'):
            return JsonResponse({
                'ok': False,
                'error': analysis_result['error']
            })
        
        # Возвращаем данные для создания объекта
        return JsonResponse({
            'ok': True,
            'object_data': analysis_result,
            'redirect_url': '/objects/create/',
            'message': 'Данные извлечены из сообщения для создания объекта'
        })
        
    except Exception as e:
        logger.error(f"Ошибка создания объекта из сообщения: {e}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def create_object_from_json(request):
    """Создание объекта в базе данных на основе JSON из последнего файла"""
    try:
        from object.models import Objekt, KategoriyaResursa, Resurs, ResursyPoObjektu, FakticheskijResursPoObjektu
        from sotrudniki.models import Specialnost, Podrazdelenie
        from datetime import date
        import os
        import glob
        from django.conf import settings
        
        # Находим последний JSON файл в папке media/documents_ai
        json_dir = os.path.join(settings.MEDIA_ROOT, 'documents_ai')
        json_files = glob.glob(os.path.join(json_dir, '*.json'))
        
        if not json_files:
            return JsonResponse({'ok': False, 'error': 'Нет JSON файлов в папке documents_ai'})
        
        # Берем последний файл по времени модификации
        latest_json_file = max(json_files, key=os.path.getmtime)
        
        # Читаем JSON данные
        with open(latest_json_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Извлекаем JSON из содержимого
        import re
        
        # Ищем JSON в markdown блоке
        json_match = re.search(r'```json\s*\n(.*?)\n```', content, re.DOTALL)
        if json_match:
            json_str = json_match.group(1).strip()
        else:
            # Ищем JSON объект по фигурным скобкам
            start = content.find('{')
            if start == -1:
                raise ValueError('JSON объект не найден в файле')
            
            brace_count = 0
            end = start
            for i, char in enumerate(content[start:], start):
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        end = i + 1
                        break
            
            json_str = content[start:end].strip()
        
        json_data = json.loads(json_str)
        
        # Создаем объект
        obj = Objekt.objects.create(
            nazvanie='Объект из AI',
            data_nachala=date.today(),
            data_plan_zaversheniya=date.today(),
            otvetstvennyj='Администратор'
        )
        
        # Привязываем объект к организациям авторизованного пользователя
        if request.user.is_authenticated:
            try:
                user_profile = request.user.profile
                user_organizations = user_profile.organizations.all()
                obj.organizacii.set(user_organizations)
            except:
                # Если профиль не найден, пропускаем
                pass
        
        # Получаем подразделение с кодом 792
        podrazdelenie, _ = Podrazdelenie.objects.get_or_create(
            kod='792',
            defaults={'nazvanie': 'Линейные сотрудники'}
        )
        
        # Обрабатываем kategoriya_resursa (расходы)
        if 'kategoriya_resursa' in json_data:
            for category_name, items in json_data['kategoriya_resursa'].items():
                # Создаем категорию расходов
                category, _ = KategoriyaResursa.objects.get_or_create(
                    nazvanie=category_name,
                    defaults={'raskhod_dokhod': True}
                )
                
                for item in items:
                    if category_name in ['Кадры', 'Кадровое обеспечение']:
                        resource_name = item.get('sotrudniki_specialnost') or item.get('name', 'Не указано')
                        quantity = item.get('hours') or item.get('часов', 1)
                        price = item.get('price') or item.get('стоимость_часа', 0)
                        unit = 'час'
                        
                        Specialnost.objects.get_or_create(
                            nazvanie=resource_name,
                            defaults={'kategoriya': 'Строительство'}
                        )
                    else:
                        resource_name = item.get('name') or item.get('наименование', 'Не указано')
                        quantity = item.get('count') or item.get('hours') or item.get('количество') or item.get('часов', 1)
                        price = item.get('price') or item.get('цена_за_ед') or item.get('стоимость_часа', 0)
                        unit = item.get('unit') or item.get('ед_изм', 'шт')
                    
                    # Создаем ресурс
                    resource, _ = Resurs.objects.get_or_create(
                        naimenovanie=resource_name,
                        kategoriya_resursa=category,
                        defaults={'edinica_izmereniya': unit}
                    )
                    
                    # Добавляем ресурс к объекту
                    resurs_po_objektu = ResursyPoObjektu.objects.create(
                        objekt=obj,
                        resurs=resource,
                        kolichestvo=quantity,
                        cena=price
                    )
                    
                    # Создаем фактический ресурс
                    FakticheskijResursPoObjektu.objects.create(
                        resurs_po_objektu=resurs_po_objektu
                    )
        
        # Обрабатываем works (доходы)
        if 'works' in json_data:
            for work_section in json_data['works']:
                section_name = work_section.get('section', 'Работы')
                
                # Создаем категорию доходов
                category, _ = KategoriyaResursa.objects.get_or_create(
                    nazvanie=section_name,
                    defaults={'raskhod_dokhod': False}
                )
                
                for item in work_section.get('items', []):
                    resource_name = item.get('name') or item.get('наименование', 'Не указано')
                    quantity = item.get('count') or item.get('количество', 1)
                    unit = item.get('unit') or item.get('ед_изм', 'шт')
                    price = item.get('price') or item.get('цена_за_ед', 0)
                    
                    # Создаем ресурс
                    resource, _ = Resurs.objects.get_or_create(
                        naimenovanie=resource_name,
                        kategoriya_resursa=category,
                        defaults={'edinica_izmereniya': unit}
                    )
                    
                    # Добавляем ресурс к объекту с ценой из JSON (доходы)
                    resurs_po_objektu = ResursyPoObjektu.objects.create(
                        objekt=obj,
                        resurs=resource,
                        kolichestvo=quantity,
                        cena=price
                    )
                    
                    # Создаем фактический ресурс
                    FakticheskijResursPoObjektu.objects.create(
                        resurs_po_objektu=resurs_po_objektu
                    )
        
        # Добавляем всех сотрудников в подразделение 792
        from sotrudniki.models import Sotrudnik
        sotrudniki_792 = Sotrudnik.objects.filter(podrazdelenie__kod='792')
        obj.sotrudniki.set(sotrudniki_792)
        
        return JsonResponse({
            'ok': True,
            'object_id': obj.id,
            'object_name': obj.nazvanie,
            'json_file': os.path.basename(latest_json_file),
            'message': f'Объект успешно создан из файла {os.path.basename(latest_json_file)}'
        })
        
    except Exception as e:
        logger.error(f'Ошибка создания объекта из JSON: {e}')
        return JsonResponse({'ok': False, 'error': str(e)})

@csrf_exempt
@require_POST
def download_and_save_document(request):
    """Скачивание документа и сохранение в ai_chatmessage"""
    try:
        data = json.loads(request.body)
        content = data.get('content', '').strip()
        file_format = data.get('format', '').lower()
        
        if not content or file_format not in ['docx', 'pdf', 'xls', 'json']:
            return JsonResponse({'error': 'Некорректные параметры'}, status=400)
        
        # Создаем файл
        filename = f"document_{timezone.now().strftime('%Y%m%d_%H%M%S')}.{file_format}"
        
        if file_format == 'docx':
            doc = Document()
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            file_content = buf.getvalue()
        elif file_format == 'pdf':
            pdf = FPDF()
            pdf.add_page()
            font_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'fonts', 'DejaVuSans.ttf')
            if os.path.exists(font_path):
                pdf.add_font('DejaVu', '', font_path, uni=True)
                pdf.set_font('DejaVu', '', 10)
            else:
                pdf.set_font('Arial', '', 10)
            
            for line in content.split('\n'):
                if line.strip():
                    try:
                        pdf.multi_cell(0, 8, line)
                    except:
                        pdf.cell(0, 8, line[:50] + '...' if len(line) > 50 else line)
                        pdf.ln()
            
            buf = io.BytesIO()
            pdf.output(buf)
            file_content = buf.getvalue()
        elif file_format == 'json':
            # Для JSON формата сохраняем как есть
            file_content = content.encode('utf-8')
        else:  # xls
            lines = [l for l in content.split('\n') if l.strip()]
            df = pd.DataFrame({'Документ': lines})
            buf = io.BytesIO()
            with pd.ExcelWriter(buf, engine='xlsxwriter') as writer:
                df.to_excel(writer, index=False)
            file_content = buf.getvalue()
        
        # Сохраняем файл в media/documents_ai/
        from django.core.files.base import ContentFile
        from django.core.files.storage import default_storage
        
        file_path = default_storage.save(f'documents_ai/{filename}', ContentFile(file_content))
        
        # Сохраняем в ai_chatmessage
        from ai.models import ChatSession, ChatMessage
        
        # Получаем или создаем сессию
        session, created = ChatSession.objects.get_or_create(
            session_id='telegram_downloads',
            defaults={'session_id': 'telegram_downloads'}
        )
        
        # Создаем запись о скачанном документе
        ChatMessage.objects.create(
            session=session,
            message_type='system',
            content=f'Скачан документ: {filename}',
            file=file_path
        )
        
        # Возвращаем файл для скачивания
        response = HttpResponse(file_content, content_type='application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        logger.error(f'Ошибка скачивания и сохранения: {e}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def save_json_response(request):
    """Сохранение JSON ответа в файл и базу данных"""
    try:
        data = json.loads(request.body)
        content = data.get('content', '')
        json_data = data.get('json_data', '')
        
        if not json_data:
            return JsonResponse({'ok': False, 'error': 'JSON данные не найдены'})
        
        from django.core.files.base import ContentFile
        from django.core.files.storage import default_storage
        import uuid
        from django.utils import timezone
        
        # Создаем уникальное имя файла
        filename = f"ai_response_{timezone.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.json"
        file_path = f'documents_ai/{filename}'
        
        # Создаем папку если не существует
        import os
        media_dir = os.path.join(settings.MEDIA_ROOT, 'documents_ai')
        os.makedirs(media_dir, exist_ok=True)
        
        # Сохраняем JSON в файл
        saved_path = default_storage.save(file_path, ContentFile(json_data.encode('utf-8')))
        logger.info(f'JSON файл сохранен: {saved_path}')
        
        # Сохраняем в таблицу ai_chatmessage
        from ai.models import ChatSession, ChatMessage
        
        session, created = ChatSession.objects.get_or_create(
            session_id='telegram_json_responses',
            defaults={'session_id': 'telegram_json_responses'}
        )
        
        ChatMessage.objects.create(
            session=session,
            message_type='assistant',
            content=content,
            file=saved_path
        )
        
        return JsonResponse({
            'ok': True,
            'file_path': saved_path,
            'message': 'JSON ответ сохранен в файл и базу данных'
        })
        
    except Exception as e:
        logger.error(f'Ошибка сохранения JSON ответа: {e}')
        return JsonResponse({'ok': False, 'error': str(e)})
@csrf_exempt
@require_POST
def save_hours(request):
    """Сохранение часов в табель и расход ресурсов"""
    try:
        data = json.loads(request.body)
        hours_data = data.get('hours', [])
        objekt_id = data.get('objekt_id')
        date = data.get('date')
        
        if not hours_data or not objekt_id or not date:
            return JsonResponse({
                'success': False,
                'error': 'Необходимо указать данные о часах, ID объекта и дату'
            })
        
        # Импортируем модели
        from sotrudniki.models import Sotrudnik, SotrudnikiZarplaty
        from object.models import Objekt, ResursyPoObjektu, Resurs, KategoriyaResursa, FakticheskijResursPoObjektu, RaskhodResursa
        from django.utils import timezone
        import datetime
        
        # Проверяем существование объекта
        try:
            objekt = Objekt.objects.get(id=objekt_id)
        except Objekt.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': f'Объект с ID {objekt_id} не найден'
            })
        
        # Преобразуем дату из строки в объект даты
        try:
            # Пробуем разные форматы даты
            date_formats = ['%d.%m.%Y', '%d.%m.%y', '%Y-%m-%d', '%d/%m/%Y', '%d/%m/%y']
            parsed_date = None
            
            for date_format in date_formats:
                try:
                    parsed_date = datetime.datetime.strptime(date, date_format).date()
                    break
                except ValueError:
                    continue
            
            if not parsed_date:
                # Если не удалось распознать дату, используем текущую
                parsed_date = timezone.now().date()
        except Exception as e:
            logger.error(f"Ошибка преобразования даты: {e}")
            parsed_date = timezone.now().date()
        
        # Получаем или создаем категорию ресурса для зарплаты
        zarplata_category, _ = KategoriyaResursa.objects.get_or_create(
            nazvanie='Зарплата',
            defaults={'raskhod_dokhod': True}  # Расход
        )
        
        saved_hours = []
        saved_resources = []
        errors = []
        
        for item in hours_data:
            try:
                employee_id = item.get('employee_id')
                employee_fio = item.get('employee_fio')
                hours = float(item.get('hours', 0))
                kpi = float(item.get('kpi', 1.0))
                position = item.get('position', '')
                
                if hours <= 0:
                    continue  # Пропускаем записи с нулевыми часами
                
                # Находим сотрудника по ID или по ФИО
                sotrudnik = None
                if employee_id:
                    try:
                        sotrudnik = Sotrudnik.objects.get(id=employee_id)
                    except Sotrudnik.DoesNotExist:
                        pass
                
                if not sotrudnik and employee_fio:
                    # Ищем по ФИО
                    sotrudnik = Sotrudnik.objects.filter(fio__icontains=employee_fio).first()
                
                if not sotrudnik:
                    errors.append(f'Сотрудник не найден: {employee_fio}')
                    continue
                
                # Создаем запись в таблице SotrudnikiZarplaty
                zarplata, created = SotrudnikiZarplaty.objects.update_or_create(
                    sotrudnik=sotrudnik,
                    objekt=objekt,
                    data=parsed_date,
                    defaults={
                        'kolichestvo_chasov': hours,
                        'kpi': kpi,
                        'vydano': False
                    }
                )
                
                saved_hours.append({
                    'id': zarplata.id,
                    'sotrudnik': sotrudnik.fio,
                    'hours': hours,
                    'created': created
                })
                
                # Создаем ресурс для зарплаты
                resource_name = f'Зарплата {sotrudnik.fio}'
                
                # Получаем или создаем ресурс
                resource, _ = Resurs.objects.get_or_create(
                    naimenovanie=resource_name,
                    kategoriya_resursa=zarplata_category,
                    defaults={'edinica_izmereniya': 'час'}
                )
                
                # Создаем запись в таблице ResursyPoObjektu
                resurs_po_objektu, created = ResursyPoObjektu.objects.update_or_create(
                    objekt=objekt,
                    resurs=resource,
                    defaults={
                        'kolichestvo': hours,
                        'cena': 0  # Цена будет установлена позже
                    }
                )
                
                # Создаем запись в таблице FakticheskijResursPoObjektu
                fakticheskij_resurs, _ = FakticheskijResursPoObjektu.objects.update_or_create(
                    resurs_po_objektu=resurs_po_objektu
                )
                
                # Создаем или обновляем запись в таблице RaskhodResursa
                # Используем update_or_create для обеспечения уникальности по fakticheskij_resurs и data
                raskhod, _ = RaskhodResursa.objects.update_or_create(
                    fakticheskij_resurs=fakticheskij_resurs,
                    data=parsed_date,
                    defaults={
                        'izraskhodovano': hours
                    }
                )
                
                saved_resources.append({
                    'id': resurs_po_objektu.id,
                    'resource': resource_name,
                    'hours': hours,
                    'created': created
                })
                
            except Exception as e:
                logger.error(f"Ошибка сохранения часов: {e}")
                errors.append(str(e))
        
        return JsonResponse({
            'success': True,
            'saved_hours': saved_hours,
            'saved_resources': saved_resources,
            'errors': errors,
            'message': f'Сохранено {len(saved_hours)} записей в табель'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Некорректный JSON'
        }, status=400)
    except Exception as e:
        logger.error(f"Ошибка сохранения часов: {e}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@csrf_exempt
def get_json_files(request):
    """Получение списка JSON файлов из папки media/documents_ai"""
    try:
        import os
        import glob
        from django.conf import settings
        
        # Путь к папке с JSON файлами
        json_dir = os.path.join(settings.MEDIA_ROOT, 'documents_ai')
        
        if not os.path.exists(json_dir):
            return JsonResponse({'ok': False, 'error': 'Папка documents_ai не найдена'})
        
        # Получаем все JSON файлы
        json_files = glob.glob(os.path.join(json_dir, '*.json'))
        
        if not json_files:
            return JsonResponse({'ok': False, 'error': 'JSON файлы не найдены'})
        
        # Формируем список файлов с информацией
        files_info = []
        for file_path in json_files:
            file_name = os.path.basename(file_path)
            file_stat = os.stat(file_path)
            
            files_info.append({
                'name': file_name,
                'size': file_stat.st_size,
                'modified': file_stat.st_mtime
            })
        
        # Сортируем по времени модификации (новые сначала)
        files_info.sort(key=lambda x: x['modified'], reverse=True)
        
        return JsonResponse({
            'ok': True,
            'files': files_info,
            'count': len(files_info)
        })
        
    except Exception as e:
        logger.error(f'Ошибка получения списка JSON файлов: {e}')
        return JsonResponse({'ok': False, 'error': str(e)})

@csrf_exempt
@require_POST
def create_object_from_selected_json(request):
    """Создание объекта в базе данных на основе выбранного JSON файла"""
    try:
        data = json.loads(request.body)
        file_name = data.get('file_name', '').strip()
        
        if not file_name:
            return JsonResponse({'ok': False, 'error': 'Не указано имя файла'})
        
        from object.models import Objekt, KategoriyaResursa, Resurs, ResursyPoObjektu, FakticheskijResursPoObjektu
        from sotrudniki.models import Specialnost, Podrazdelenie
        from datetime import date
        import os
        from django.conf import settings
        
        # Путь к выбранному JSON файлу
        json_file_path = os.path.join(settings.MEDIA_ROOT, 'documents_ai', file_name)
        
        if not os.path.exists(json_file_path):
            return JsonResponse({'ok': False, 'error': f'Файл {file_name} не найден'})
        
        # Читаем JSON данные
        with open(json_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Извлекаем JSON из содержимого
        import re
        
        # Ищем JSON в markdown блоке
        json_match = re.search(r'```json\s*\n(.*?)\n```', content, re.DOTALL)
        if json_match:
            json_str = json_match.group(1).strip()
        else:
            # Ищем JSON объект по фигурным скобкам
            start = content.find('{')
            if start == -1:
                raise ValueError('JSON объект не найден в файле')
            
            brace_count = 0
            end = start
            for i, char in enumerate(content[start:], start):
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        end = i + 1
                        break
            
            json_str = content[start:end].strip()
        
        json_data = json.loads(json_str)
        
        # Создаем объект с именем на основе файла
        object_name = f'Объект из {file_name}'
        obj = Objekt.objects.create(
            nazvanie=object_name,
            data_nachala=date.today(),
            data_plan_zaversheniya=date.today(),
            otvetstvennyj='Администратор'
        )
        
        # Привязываем объект к организациям авторизованного пользователя
        if request.user.is_authenticated:
            try:
                user_profile = request.user.profile
                user_organizations = user_profile.organizations.all()
                obj.organizacii.set(user_organizations)
            except:
                # Если профиль не найден, пропускаем
                pass
        
        # Получаем подразделение с кодом 792
        podrazdelenie, _ = Podrazdelenie.objects.get_or_create(
            kod='792',
            defaults={'nazvanie': 'Линейные сотрудники'}
        )
        
        # Обрабатываем kategoriya_resursa (расходы)
        if 'kategoriya_resursa' in json_data:
            for category_name, items in json_data['kategoriya_resursa'].items():
                # Создаем категорию расходов
                category, _ = KategoriyaResursa.objects.get_or_create(
                    nazvanie=category_name,
                    defaults={'raskhod_dokhod': True}
                )
                
                for item in items:
                    if category_name in ['Кадры', 'Кадровое обеспечение']:
                        resource_name = item.get('sotrudniki_specialnost') or item.get('name', 'Не указано')
                        quantity = item.get('hours') or item.get('часов', 1)
                        price = item.get('price') or item.get('стоимость_часа', 0)
                        unit = 'час'
                        
                        Specialnost.objects.get_or_create(
                            nazvanie=resource_name,
                            defaults={'kategoriya': 'Строительство'}
                        )
                    else:
                        resource_name = item.get('name') or item.get('наименование', 'Не указано')
                        quantity = item.get('count') or item.get('hours') or item.get('количество') or item.get('часов', 1)
                        price = item.get('price') or item.get('цена_за_ед') or item.get('стоимость_часа', 0)
                        unit = item.get('unit') or item.get('ед_изм', 'шт')
                    
                    # Создаем ресурс
                    resource, _ = Resurs.objects.get_or_create(
                        naimenovanie=resource_name,
                        kategoriya_resursa=category,
                        defaults={'edinica_izmereniya': unit}
                    )
                    
                    # Добавляем ресурс к объекту
                    resurs_po_objektu = ResursyPoObjektu.objects.create(
                        objekt=obj,
                        resurs=resource,
                        kolichestvo=quantity,
                        cena=price
                    )
                    
                    # Создаем фактический ресурс
                    FakticheskijResursPoObjektu.objects.create(
                        resurs_po_objektu=resurs_po_objektu
                    )
        
        # Обрабатываем works (доходы)
        if 'works' in json_data:
            for work_section in json_data['works']:
                section_name = work_section.get('section', 'Работы')
                
                # Создаем категорию доходов
                category, _ = KategoriyaResursa.objects.get_or_create(
                    nazvanie=section_name,
                    defaults={'raskhod_dokhod': False}
                )
                
                for item in work_section.get('items', []):
                    resource_name = item.get('name') or item.get('наименование', 'Не указано')
                    quantity = item.get('count') or item.get('количество', 1)
                    unit = item.get('unit') or item.get('ед_изм', 'шт')
                    price = item.get('price') or item.get('цена_за_ед', 0)
                    
                    # Создаем ресурс
                    resource, _ = Resurs.objects.get_or_create(
                        naimenovanie=resource_name,
                        kategoriya_resursa=category,
                        defaults={'edinica_izmereniya': unit}
                    )
                    
                    # Добавляем ресурс к объекту с ценой из JSON (доходы)
                    resurs_po_objektu = ResursyPoObjektu.objects.create(
                        objekt=obj,
                        resurs=resource,
                        kolichestvo=quantity,
                        cena=price
                    )
                    
                    # Создаем фактический ресурс
                    FakticheskijResursPoObjektu.objects.create(
                        resurs_po_objektu=resurs_po_objektu
                    )
        
        # Добавляем всех сотрудников в подразделение 792
        from sotrudniki.models import Sotrudnik
        sotrudniki_792 = Sotrudnik.objects.filter(podrazdelenie__kod='792')
        obj.sotrudniki.set(sotrudniki_792)
        
        return JsonResponse({
            'ok': True,
            'object_id': obj.id,
            'object_name': obj.nazvanie,
            'json_file': file_name,
            'message': f'Объект успешно создан из файла {file_name}'
        })
        
    except Exception as e:
        logger.error(f'Ошибка создания объекта из выбранного JSON: {e}')
        return JsonResponse({'ok': False, 'error': str(e)})

@csrf_exempt
@require_POST
def find_employees(request):
    """Поиск сотрудников по фамилии и другим словам в сообщении, похожим на ФИО в таблице sotrudniki_sotrudnik"""
    try:
        data = json.loads(request.body)
        surnames = data.get('surnames', [])
        objekt_id = data.get('objekt_id')
        
        if not surnames or not objekt_id:
            return JsonResponse({
                'success': False,
                'error': 'Необходимо указать фамилии и ID объекта'
            })
        
        logger.info(f"Поиск сотрудников по словам: {surnames} на объекте {objekt_id}")
        
        # Импортируем модели
        from sotrudniki.models import Sotrudnik
        from object.models import Objekt
        
        # Проверяем существование объекта
        try:
            objekt = Objekt.objects.get(id=objekt_id)
        except Objekt.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': f'Объект с ID {objekt_id} не найден'
            })
        
        # Получаем сотрудников, связанных с объектом
        objekt_sotrudniki = objekt.sotrudniki.all()
        
        # Если нет сотрудников на объекте, возвращаем пустой результат
        if not objekt_sotrudniki.exists():
            return JsonResponse({
                'success': True,
                'employees': [],
                'not_found': surnames,
                'message': f'На объекте "{objekt.nazvanie}" нет сотрудников'
            })
        
        # Получаем всех сотрудников объекта
        all_employees = []
        for sotrudnik in objekt_sotrudniki:
            
            # Получаем специальность сотрудника
            specialnost = ""
            if hasattr(sotrudnik, 'specialnost') and sotrudnik.specialnost:
                specialnost = sotrudnik.specialnost.nazvanie
            
            # Извлекаем части ФИО
            fio_parts = sotrudnik.fio.split()
            surname = fio_parts[0] if fio_parts else ""
            
            # Добавляем все части ФИО для поиска
            name_parts = []
            if len(fio_parts) > 0:
                name_parts.append(fio_parts[0])  # Фамилия
            if len(fio_parts) > 1:
                name_parts.append(fio_parts[1])  # Имя
            if len(fio_parts) > 2:
                name_parts.append(fio_parts[2])  # Отчество
            
            all_employees.append({
                'id': sotrudnik.id,
                'fio': sotrudnik.fio,
                'surname': surname,
                'name_parts': name_parts,
                'specialnost': specialnost
            })
        
        # Функция для сравнения слов с учетом возможных ошибок
        def is_similar_word(word1, word2):
            # Приводим к нижнему регистру и удаляем лишние символы
            clean1 = word1.lower().replace(' ', '').replace('-', '')
            clean2 = word2.lower().replace(' ', '').replace('-', '')
            
            # Точное совпадение
            if clean1 == clean2:
                return True
            
            # Если одно слово содержит другое и длиннее не более чем на 3 символа
            if clean1 in clean2 and len(clean2) - len(clean1) <= 3:
                return True
            if clean2 in clean1 and len(clean1) - len(clean2) <= 3:
                return True
            
            # Реализуем свою функцию расстояния Левенштейна
            def levenshtein_distance(s1, s2):
                if len(s1) < len(s2):
                    return levenshtein_distance(s2, s1)
                if len(s2) == 0:
                    return len(s1)
                
                previous_row = range(len(s2) + 1)
                for i, c1 in enumerate(s1):
                    current_row = [i + 1]
                    for j, c2 in enumerate(s2):
                        insertions = previous_row[j + 1] + 1
                        deletions = current_row[j] + 1
                        substitutions = previous_row[j] + (c1 != c2)
                        current_row.append(min(insertions, deletions, substitutions))
                    previous_row = current_row
                
                return previous_row[-1]
            
            # Для коротких слов допускаем только 1 ошибку
            if len(clean1) <= 5 and len(clean2) <= 5:
                return levenshtein_distance(clean1, clean2) <= 1
            
            # Для более длинных слов допускаем до 2 ошибок
            return levenshtein_distance(clean1, clean2) <= 2
        
        # Ищем совпадения
        found_employees = []
        not_found_surnames = []
        found_employee_ids = set()  # Множество для отслеживания найденных ID сотрудников
        processed_search_words = set()  # Множество для отслеживания обработанных поисковых слов
        
        for search_word in surnames:
            # Пропускаем поисковое слово, если оно уже было обработано
            if search_word.lower() in processed_search_words:
                continue
                
            processed_search_words.add(search_word.lower())  # Добавляем слово в обработанные
            found = False
            
            for employee in all_employees:
                # Пропускаем сотрудника, если он уже был найден
                if employee['id'] in found_employee_ids:
                    continue
                    
                # Проверяем совпадение с фамилией
                if is_similar_word(search_word, employee['surname']):
                    found_employees.append({
                        'id': employee['id'],
                        'fio': employee['fio'],
                        'specialnost': employee['specialnost'],
                        'matched_word': search_word,
                        'match_type': 'фамилия'
                    })
                    found_employee_ids.add(employee['id'])  # Добавляем ID в множество найденных
                    found = True
                    break
                
                # Проверяем совпадение с любой частью ФИО
                for i, name_part in enumerate(employee['name_parts']):
                    if is_similar_word(search_word, name_part):
                        match_type = 'фамилия' if i == 0 else ('имя' if i == 1 else 'отчество')
                        found_employees.append({
                            'id': employee['id'],
                            'fio': employee['fio'],
                            'specialnost': employee['specialnost'],
                            'matched_word': search_word,
                            'match_type': match_type
                        })
                        found_employee_ids.add(employee['id'])  # Добавляем ID в множество найденных
                        found = True
                        break
                
                # Проверяем совпадение с полным ФИО
                if not found and is_similar_word(search_word, employee['fio']):
                    found_employees.append({
                        'id': employee['id'],
                        'fio': employee['fio'],
                        'specialnost': employee['specialnost'],
                        'matched_word': search_word,
                        'match_type': 'полное ФИО'
                    })
                    found_employee_ids.add(employee['id'])  # Добавляем ID в множество найденных
                    found = True
                    break
                
                if found:
                    break
            
            if not found:
                not_found_surnames.append(search_word)
        
        # Проходим по всем фамилиям еще раз, чтобы найти совпадения для оставшихся слов
        # с сотрудниками, которые еще не были найдены
        for search_word in surnames:
            # Пропускаем слово, если оно уже было обработано и найдено
            if search_word in not_found_surnames:
                found = False
                
                for employee in all_employees:
                    # Пропускаем сотрудника, если он уже был найден
                    if employee['id'] in found_employee_ids:
                        continue
                        
                    # Проверяем совпадение с фамилией
                    if is_similar_word(search_word, employee['surname']):
                        found_employees.append({
                            'id': employee['id'],
                            'fio': employee['fio'],
                            'specialnost': employee['specialnost'],
                            'matched_word': search_word,
                            'match_type': 'фамилия (доп. поиск)'
                        })
                        found_employee_ids.add(employee['id'])  # Добавляем ID в множество найденных
                        found = True
                        break
                
                if found:
                    # Удаляем из списка ненайденных
                    not_found_surnames.remove(search_word)
        
        return JsonResponse({
            'success': True,
            'employees': found_employees,
            'not_found': not_found_surnames,
            'message': f'Найдено {len(found_employees)} сотрудников из {len(set(surnames))} запрошенных слов'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Некорректный JSON'
        }, status=400)
    except Exception as e:
        logger.error(f"Ошибка поиска сотрудников: {e}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)